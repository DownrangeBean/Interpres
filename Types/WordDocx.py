import os
import logging
from docx import Document
from Types.TranslatableDocument import TranslatableDocument
from Util.Logging import get_logger
import Interpres_Globals

logger = get_logger(__name__)
logger.setLevel(Interpres_Globals.VERBOSITY)


class WordDocx(TranslatableDocument):

    def __init__(self, filepath):
        super(WordDocx, self).__init__(filepath)
        self._new_doc

    def _openDoc(self):
        # print(os.path.join(self.dirpath, '.'.join((self.base, self.ext))))
        self.document = Document(os.path.join(self.dirpath, '.'.join((self.base, self.ext))))

    def save(self, name=None):
        if len(self.destination_text) > 0:
            i = 0
            for para in self._new_doc.paragraphs:
                for run in para.runs:
                    run.text = self.destination_text[i]
                    logger.debug('destination text: "%s"', self.destination_text[i])
                    logger.debug('same text but in run: "%s"', run.text)
                    i += 1
                logger.debug('text in paragraph: %s', para.text)

            for table in self.document.tables:
                for row in table.rows:
                    for i in range(len(row.cells)):
                        row.cells[i].text = self.destination_text[i]
            self._new_doc.save(self._checkdir(name))
        else:
            logger.error("Unable to save as no data has been parsed yet.")

    def _extract_text(self):
        self._new_doc = Document()  # self.dirpath + self.base + '_' + 'tmp' + self.ext
        for para in self.document.paragraphs:
            new_para = self._new_doc.add_paragraph()
            p = Paragraph(para)
            for run in p.runs:
                r = new_para.add_run(run.text)
                r.bold = run.bold
                r.italic = run.italic
                r.underline = run.underline
                r.style = run.style
                self.source_text.append(run.text)
                logger.debug('extracting run:' + run.text)

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.source_text.append(cell.text)


class Paragraph(object):

    def __init__(self, para):
        self._para = para
        self._runs = []
        self._text = ''

    @property
    def text(self):
        if not self._text:
            if len(self.runs) > 1:
                self._text = ''.join([run.text for run in self.runs])
            elif len(self.runs) > 0:
                self._text = self.runs[0].text
            else:
                self._text = ''
        logger.debug('text getter: ' + self._text)
        return self._text

    @text.setter
    def text(self, text):
        self._text = text

    @text.deleter
    def text(self):
        self._text = ''

    @property
    def runs(self):
        if self._runs:
            return self._runs

        page = None
        for j in range(len(self._para.runs)):
            cond = False
            if page:
                cond = page.text[-1] == ' ' or page.text[-1] == '\t' or page.text[-1] == '\n'
                logger.debug('using page condition')
            else:
                logger.debug('using run condition')
                if self._para.runs[j].text:
                    cond = self._para.runs[j].text[-1] == ' ' or self._para.runs[j].text[-1] == '\t' or \
                           self._para.runs[j].text[-1] == '\n'
            if cond:
                if page:
                    logger.debug('page: ' + str(page.__dict__))
                    self._runs.append(page)
                run = RunMessenger(paragraph=j, text=self._para.runs[j].text
                                                , bold=self._para.runs[j].bold
                                                , italic=self._para.runs[j].italic
                                                , underline=self._para.runs[j].underline
                                                , style=self._para.runs[j].style)
                self._runs.append(run)
                logger.debug('run: ' + str(run.__dict__))
                page = None

            elif page is None:
                page = RunMessenger(paragraph=j, text=self._para.runs[j].text
                               , bold=self._para.runs[j].bold
                               , italic=self._para.runs[j].italic
                               , underline=self._para.runs[j].underline
                               , style=self._para.runs[j].style)
            else:
                page.text += self._para.runs[j].text

        if page:
            logger.debug('page: ' + str(page.__dict__))
            self._runs.append(page)
        return self._runs

    @runs.setter
    def runs(self, value):
        self._runs = value


class RunMessenger:
    def __init__(self, **kwargs):
        self.__dict__ = kwargs



