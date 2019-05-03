import unittest, os, logging, sys
from Types.WordDocx import WordDocx
from TranslationTools.Translator import Translator
from docx import Document
from Tests.WarningDecorators import ignore_warnings
import Definitions


source_text_for_test = list()
destination_text_for_test = list()


class TestTextExtraction(unittest.TestCase):

    @staticmethod
    def comp_docx(f1, f2):
        p1 = f1.paragraphs

        p2 = f2.paragraphs
        for i in range(len(p1)):
            if p1[i]._line == p2[i]._line:
                continue
            else:
                print(p1[i]._line, p2[i]._line)
                return False
        return True

    @ignore_warnings
    def test_file_output(self):
        document = WordDocx(os.path.join(TestTextExtraction.filepath, TestTextExtraction.filename))
        name_extra = 'output_test'
        tx = Translator(source='cs', destination='cs', extra=name_extra)
        try:
            tx(document)
        except ValueError as e:
            logging.error('Google has blocked your current IP. Testing not possible.')

        # Save file
        document.save()

        f1 = Document(os.path.join(TestTextExtraction.filepath, TestTextExtraction.filename))
        f2 = Document(os.path.join(TestTextExtraction.filepath, TestTextExtraction.filename.split('.')[0] + '_'
                                   + name_extra + '.' + TestTextExtraction.filename.split('.')[1]))
        self.assertTrue(TestTextExtraction.comp_docx(f1, f2), 'File output is not as expected.')

#    @ignore_warnings
#    def test_extraction_method(self):
#            self.assertEqual(TestTextExtraction.document.source_text, source_text_for_test)

    @ignore_warnings
    def test_tx_method(self):
            assertcount = 0
            for i in range(len(TestTextExtraction.document.destination_text)):
                try:
                    self.assertEqual(TestTextExtraction.document.destination_text[i].strip(), destination_text_for_test[i].strip())
                except AssertionError:
                    print('index: ', i)
                    assertcount += 1

            self.assertLessEqual(assertcount,len(destination_text_for_test) / 2)

    @classmethod
    @ignore_warnings
    def tearDownClass(cls):
        cls.document.save()

    @classmethod
    @ignore_warnings
    def setUpClass(cls):
        cls.filepath = os.path.join(Definitions.TEST_OUT_DIR, 'Word', 'IODataCorrelation')
        cls.filename = 'TestDoc.docx'

        document = Document()
        document.add_heading(source_text_for_test[0])
        document.add_paragraph(source_text_for_test[1])

        heading = document.add_heading(level=2)
        heading.add_run(source_text_for_test[2])
        bold_run1 = heading.add_run(source_text_for_test[3])
        bold_run1.bold
        heading.add_run(source_text_for_test[4])

        bullet1 = document.add_paragraph(style='List Bullet')
        bold_run1 = bullet1.add_run(source_text_for_test[5])
        bold_run1.bold = True
        bullet1.add_run(source_text_for_test[6])

        bullet2 = document.add_paragraph(style='ListBullet')
        bold_run1 = bullet2.add_run(source_text_for_test[7])
        bold_run1.bold = True
        bullet2.add_run(source_text_for_test[8])

        bullet3 = document.add_paragraph(style='ListBullet')
        bold_run1 = bullet3.add_run(source_text_for_test[9])
        bold_run1.bold = True
        bullet3.add_run(source_text_for_test[10])

        document.add_heading(source_text_for_test[11], level=2)
        document.add_paragraph(source_text_for_test[12])
        document.add_heading(source_text_for_test[13], level=2)
        document.add_paragraph(source_text_for_test[14])
        document.add_heading(source_text_for_test[15], level=2)
        document.add_paragraph(source_text_for_test[16])
        document.add_paragraph(source_text_for_test[17])
        try:
            document.save(os.path.join(cls.filepath, cls.filename))
        except PermissionError:
            base, ext = cls.filename.split('.')
            i = 1
            while i < 100:
                name = base + str(i) + '.' + ext
                if not os.path.exists(os.path.join(cls.filepath, name)):
                    cls.filename = name
                    break
                i += 1
            logging.info('User warning: TestDoc.docx is open in another program. '
                  'Document has been saved to {}'.format(cls.filename))
            try:
                document.save(os.path.join(cls.filepath, cls.filename))
            except PermissionError:
                logging.error('Permission Error: Failed to save document')

        cls.document = WordDocx(os.path.join(TestTextExtraction.filepath, TestTextExtraction.filename))
        cls.tx = Translator(source='cs', destination='en', extra='Extraction')
        try:
            cls.tx(cls.document)
        except ValueError as e:
            logging.error('Google has blocked your current IP. Testing not possible.')
            sys.exit()


source_text_for_test.append('Porada operátorů')
source_text_for_test.append('Deník provozu/ the logbook '
            'Zapisování hodnot nastavených a ne odečítaných ' 
            'Zapisovat až po ustálení parametrů. '
            'Jednoznačný soubor hodnot k jednomu nastavení.')
source_text_for_test.append('Podivejte se na ')
source_text_for_test.append('vdg-nas/Tandetron-manualy, ')
source_text_for_test.append('nově adresaře:')
source_text_for_test.append('manuály Tandetron 2019...')
source_text_for_test.append(' aktualizované manuály pro Tandetron.')
source_text_for_test.append('manuály pro trasy Tandetronu ')
source_text_for_test.append('…aktualizované manuály, včetně schemat tras a vakua.')
source_text_for_test.append('Rozvody medií')
source_text_for_test.append(' … schemata a popisy')
source_text_for_test.append('Manuály Tandetronu- kontrolovat svůj postup podle manuálu!')
source_text_for_test.append('U Duoplasmatronu nižší proudy Filamentu Li kanal- nahřátí n 580 trvá 3 hod., nastavovat'
            ' na 4:00 Výměna plynů u 358-2: dodržovat postup podle manualu, tj.den předem čerpat '
            'trubičky. 358-1 momentální kondice, výhled - apertury, civka magnetu zdroje Targety '
            'Sputteru – zapisovat')
source_text_for_test.append('Manuály iontových tras- kontrolovat svůj postup podle manuálu!')
source_text_for_test.append('Před implantací SW provést magnet Degauss!! Nově postupy na dovedení svazku do komory. '
            'seznamit se s ovladaním slitů-používat.')
source_text_for_test.append('V terminu 16-17.5.:')
source_text_for_test.append('Opakovací školení ovládání vakaua , povede Pavel Plocek. '
            'Provedeme údržbové práce a inventuru nahradních dílů.')
source_text_for_test.append('Plán údržby Další návrhy')

destination_text_for_test.append("Operators meeting")
destination_text_for_test.append("Logbook / logbook Writing the values ​​set and not read Write until the parameters have stabilized. Unique set of values ​​for one setting.")
destination_text_for_test.append("Look at ")
destination_text_for_test.append("vdg-nas / Tandetron-manuals ")
destination_text_for_test.append("new address:")
destination_text_for_test.append("manuals Tandetron 2019 ... updated manuals for Tandetron.")
destination_text_for_test.append("manuals for Tandetron routes ")
destination_text_for_test.append("… Updated manuals, including route and vacuum diagrams.")
destination_text_for_test.append("Divorces of media… schemes and descriptions")
destination_text_for_test.append("Tandetron manuals - check your manual!")
destination_text_for_test.append("For Duoplasmatron lower currents of Filament Li channel- heating n 580 takes 3 hours, set at 4:00 Gas exchange at 358-2: follow the manual procedure, ie the pre-pumping time. 358-1 Current Condition, Perspective - Apertures, Magnet Reel Source Target Sputter - Write")

destination_text_for_test.append("Ion Route Manuals - Check Your Manual!")
destination_text_for_test.append("Perform Degauss magnets before SW implantation! New procedures to bring the bundle into the chamber. Familiarize yourself with alloy handling-use.")
destination_text_for_test.append("In the period 16-17.5 .:")
destination_text_for_test.append("Refresher training of the Vacau ​​control, led by Pavel Plocek. Perform maintenance work and inventory of replacement parts.")
destination_text_for_test.append("Maintenance plan Other suggestions")

if __name__ == '__main__':
    unittest.main()
